from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.hashers import make_password
from django.contrib.auth.decorators import login_required
from django.db import IntegrityError
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.conf import settings
from django.core.mail import send_mail
from .utils import fetch_youtube_channel_data, fetch_instagram_profile
from .models import *
from .forms import ChatMessagesCreateform
from django.http import JsonResponse, HttpResponse
from django.urls import reverse

from datetime import datetime, timedelta
from django.utils import timezone

import string
import random
import io
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

from decouple import config
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from django.views.decorators.http import require_POST

def index(request):
    return render(request, 'index.html')

#############-sign up-##################
def signUp(request):
    if request.method == 'POST':
        try:
            first_name = request.POST.get('first_name')
            last_name = request.POST.get('last_name')
            username = request.POST.get('username')
            email = request.POST.get('email')
            phone_number = request.POST.get('phoneNumber')
            password = request.POST.get('password')
            account_type = request.POST.get('type')

            if User.objects.filter(username=username).exists():
                messages.error(request, "User exists. Login")
                return render(request, '404.html')
            elif User.objects.filter(email=email).exists():
                messages.error(request, "Email exists.")
                return render(request, '404.html')
            new_user = User(
                first_name=first_name,
                last_name=last_name,
                username=username,
                email=email,
                phone_number=phone_number,
                password=make_password(password), 
                account_type=account_type,
            )
            new_user.save()
            context = {'username': username, 'account_type': account_type}
            return render(request, 'platform.html', context)
        except Exception as e:
            messages.error(request, f"Error during signup: {str(e)}")
            return render(request, '404.html')
    return render(request, 'signup.html')

####################################################
def resertPassword(request):
    if request.method == 'POST':
        try:
            email = request.POST.get('email')
            user = User.objects.filter(email=email).first()
            if user:
                token = 'generate_your_token_here'
                user.password_reset_token = token
                user.save()
                return render(request, 'email_error_page.html')
            else:
                return render(request, 'email_error_page.html')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, 'email_error_page.html')
    return render(request,'resert_password.html')

####################################################
def newPassword(request):
    return render(request,'new_password.html')

####################################################
def resertDone(request):
    return render(request,'resert_done.html')

#############-Create Admin Profile-##################
def adminProfile(request):
    if request.method == 'POST':
        try:
            first_name = request.POST.get('first_name')
            last_name = request.POST.get('last_name')
            username = request.POST.get('username')
            email = request.POST.get('email')
            phone_number = request.POST.get('phoneNumber')
            password = request.POST.get('password')

            if User.objects.filter(username=username).exists():
                messages.error(request, "User exists. Login")
                return render(request, 'signin.html')

            new_user = User(
                first_name=first_name,
                last_name=last_name,
                username=username,
                email=email,
                phone_number=phone_number,
                password=make_password(password), 
                is_admin=True
            )
            new_user.save()
            Admin.objects.create(user=new_user) 
            messages.success(request,'You have successfully created an account. Login')
            return render(request, 'signin.html')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, 'signin.html')
    return render(request, 'admin_profile.html')

#############-Sponser##################
def sponsorPlatform(request):
    if request.method == 'POST':
        try:
            content_category = request.POST.get('content_category')
            website = request.POST.get('website')
            username = request.POST.get('username')

            user = User.objects.get(username=username)
            new_sponsor = Sponsor(
                user=user,
                content_category=content_category,
                website=website,
            )
            new_sponsor.save()
            messages.success(request, "Your account is set up successfully. Login")
            return redirect('signin')
        except User.DoesNotExist:
            messages.error(request, "User does not exist.")
            return redirect('signup')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, 'signin.html')
    else:
        return render(request, 'signin.html')

#############-Influecer-##################
def influencerPlatform(request):
    if request.method == 'POST':
        try:
            youtube_url = request.POST.get('youtube_url')
            instagram_url = request.POST.get('instagram_url')
            content_category = request.POST.get('content_category')
            username = request.POST.get('username')

            user = User.objects.get(username=username)
            new_influencer = Influencer(
                user=user,
                content_category=content_category,
            )
            new_influencer.save()

            if youtube_url and youtube_url.strip():
                try:
                    Platform.objects.create(influencer=new_influencer, platform_name="YouTube", platform_url=youtube_url.strip())
                except IntegrityError:
                    return render(request, 'platform_error.html')
            if instagram_url and instagram_url.strip():
                try:
                    Platform.objects.create(influencer=new_influencer, platform_name="Instagram", platform_url=instagram_url.strip())
                except IntegrityError:
                    return render(request, 'platform_error.html')
            messages.success(request, "Your account is active. Login")
            return redirect('signin')
        except User.DoesNotExist:
            messages.error(request, "User does not exist.")
            return redirect('signup')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, 'signin.html')
    else:
        return render(request, 'signin.html')

#############-log in-##################
def signIn(request):
    if request.method == 'POST':
        try:
            username = request.POST.get('username')
            password = request.POST.get('password')  
            if username and password:
                user = authenticate(username=username, password=password)
                if user:
                    login(request, user)
                    if user.is_admin:
                        return redirect('admin_home')
                    else:
                        return redirect('home')
                else:
                    messages.error(request, "Invalid credentials.")
                    return render(request, 'signin.html')
            else:
                messages.error(request, "Please provide both username and password.")
                return render(request, 'signin.html')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, 'signin.html')
    else:
        return render(request, 'signin.html')

#############-log out-##################
def signOut(request):
    logout(request)
    messages.success(request,'You have logged out')
    return redirect('index')

def sponser(request):
    return render(request, 'sponser.html')
    
#############-home-##################
@login_required(login_url='index')
def home(request):
    try:
        platforms = Platform.objects.select_related('influencer__user').all()
        data = []
        for platform in platforms:
            influencer = platform.influencer
            influencer_user = influencer.user
            influencer_data = {
                'influencer_id': influencer.influencer_id,
                'influencer_username': influencer_user.username,
                'platform_name': platform.platform_name,
                'platform_url': platform.platform_url,
                'likes': platform.likes,
                'views': platform.views,
                'subscribers': platform.subscribers,
                'comments': platform.comments,
                'videos': platform.videos,
            }
            data.append(influencer_data)
        context = {'platform_data': data}
        return render(request, 'home.html', context)
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, '404.html')

#############-Remove payment system-##################
# All payment-related views and logic have been removed.

#############-Generate payment report-##################
def process_payments():
    # Payment system removed, so return empty sets
    return set(), set()

def generate_pdf_report(request):
    paid_usernames, unpaid_usernames = process_payments()
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.drawString(100, 750, "User Payment Report")
    p.drawString(100, 730, "Paid Users:")
    y = 710
    for username in paid_usernames:
        p.drawString(100, y, username)
        y -= 20
    p.drawString(100, y - 20, "Unpaid Users:")
    y -= 40
    for username in unpaid_usernames:
        p.drawString(100, y, username)
        y -= 20
    p.showPage()
    p.save()
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="user_payment_report.pdf"'
    return response

def generate_word_report(request):
    paid_usernames, unpaid_usernames = process_payments()
    doc = Document()
    doc.add_heading('User Payment Report', 0)
    doc.add_heading('Paid Users:', level=1)
    for username in paid_usernames:
        doc.add_paragraph(username)
    doc.add_heading('Unpaid Users:', level=1)
    for username in unpaid_usernames:
        doc.add_paragraph(username)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=user_payment_report.docx'
    return response

def generate_excel_report(request):
    paid_usernames, unpaid_usernames = process_payments()
    df_paid = pd.DataFrame(list(paid_usernames), columns=['Paid Users'])
    df_unpaid = pd.DataFrame(list(unpaid_usernames), columns=['Unpaid Users'])
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df_paid.to_excel(writer, sheet_name='Paid Users', index=False)
        df_unpaid.to_excel(writer, sheet_name='Unpaid Users', index=False)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=user_payment_report.xlsx'
    return response

#############-Admin home-##################
@login_required(login_url='index')
def admin_home(request):
    try:
        # Example: show all users
        users = User.objects.all()
        context = {'users': users}
        return render(request, 'admin_home.html', context)
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, '404.html')

#############-Profile-##################
@login_required(login_url='index')
def profile(request):
    try:
        return render(request, 'profile.html')
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, '404.html')

#############-Profile-##################
@login_required(login_url='index')
def update_profile(request):
    try:
        if request.method == 'POST':
            user = request.user
            user.username = request.POST.get('username', user.username)
            user.email = request.POST.get('email', user.email)
            user.first_name = request.POST.get('first_name', user.first_name)
            user.last_name = request.POST.get('last_name', user.last_name)
            user.phone_number = request.POST.get('phone_number', user.phone_number)
            user.save()
            messages.success(request, "Profile updated successfully.")
            return redirect('update_profile')
        return render(request, 'update_profile.html')
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, 'update_profile.html')

#############-Delete account-##################
def delete_account(request):
    if request.method == 'GET':
        try:
            return render(request, 'delete_account.html')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return render(request, '404.html')
    else:
        try:
            user = request.user
            user.delete()
            messages.success(request, "Account deleted successfully.")
            return redirect('index')
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect('update_profile')

#############-Chat and messaging views-##################
def chat_room(request, room_name):
    return render(request, 'chat.html', {
        'room_name': room_name
    })

@login_required
def create_room(request):
    try:
        if request.method == 'POST':
            influencer_id = request.POST.get('influencer_id')
            sponsor_id = request.POST.get('sponsor_id')
            influencer = Influencer.objects.get(pk=influencer_id)
            sponsor = Sponsor.objects.get(pk=sponsor_id)
            room, created = Room.objects.get_or_create(influencer=influencer, sponsor=sponsor)
            return redirect('chat_room', room_id=room.room_id)
        return redirect('home')
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return redirect('home')

@login_required
@require_POST
def accept_room(request, room_id):
    try:
        room = get_object_or_404(Room, pk=room_id)
        room.accepted = True
        room.save()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
@require_POST
def send_message(request):
    try:
        data = json.loads(request.body)
        room_id = data.get('room_id')
        content = data.get('content')
        room = get_object_or_404(Room, pk=room_id)
        message = Message.objects.create(user=request.user, room=room, content=content)
        return JsonResponse({'success': True, 'sender': request.user.username, 'message': content, 'timestamp': message.timestamp})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def get_messages(request, room_id):
    try:
        room = get_object_or_404(Room, pk=room_id)
        messages_qs = Message.objects.filter(room=room).order_by('timestamp')
        messages_data = [
            {
                'user': msg.user.username if msg.user else '',
                'content': msg.content,
                'timestamp': msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
            }
            for msg in messages_qs
        ]
        return JsonResponse(messages_data, safe=False)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def chat_rooms(request):
    try:
        user = request.user
        influencer_rooms = Room.objects.filter(influencer__user=user)
        sponsor_rooms = Room.objects.filter(sponsor__user=user)
        context = {
            'influencer_rooms': influencer_rooms,
            'sponsor_rooms': sponsor_rooms,
        }
        return render(request, 'chat_rooms.html', context)
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, '404.html')

def my_view(request):
    return render(request, '404.html')

def chat(request):
    try:
        user = request.user
        if user.account_type == "INFLUENCER":
            rooms = Room.objects.filter(influencer__user=user)
        else:
            rooms = Room.objects.filter(sponsor__user=user)
        room_messages = {}
        for room in rooms:
            room_messages[room.room_id] = Message.objects.filter(room=room).order_by('timestamp')
        context = {
            'rooms': rooms,
            'room_messages': room_messages,
        }
        return render(request, 'messages.html', context)
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return render(request, '404.html')

def fetch_messages(request):
    try:
        room_id = request.GET.get('room_id')
        room = get_object_or_404(Room, pk=room_id)
        messages_qs = Message.objects.filter(room=room).order_by('timestamp')
        message_data = []
        for message in messages_qs:
            message_data.append({
                'user': message.user.username if message.user else '',
                'content': message.content,
                'timestamp': message.timestamp.strftime('%Y-%m-%d %H:%M:%S')
            })
        return JsonResponse(message_data, safe=False)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

#########Content Performance Report############
def generate_content_performance_report_excel(request):
    try:
        platforms = Platform.objects.all().order_by('-likes')
        data = []
        for platform in platforms:
            data.append({
                'Influencer': platform.influencer.user.username,
                'Platform': platform.platform_name,
                'Likes': platform.likes,
                'Views': platform.views,
                'Subscribers': platform.subscribers,
                'Comments': platform.comments,
                'Videos': platform.videos,
            })
        df = pd.DataFrame(data)
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=content_performance_report.xlsx'
        return response
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return redirect('admin_home')

def generate_content_performance_report_pdf(request):
    try:
        platforms = Platform.objects.all().order_by('-likes')
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.drawString(100, 750, "Content Performance Report")
        y = 730
        for platform in platforms:
            p.drawString(100, y, f"{platform.influencer.user.username} - {platform.platform_name}: Likes={platform.likes}, Views={platform.views}")
            y -= 20
            if y < 100:
                p.showPage()
                y = 750
        p.showPage()
        p.save()
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="content_performance_report.pdf"'
        return response
    except Exception as e:
        messages.error(request, f"Error: {str(e)}")
        return redirect('admin_home')

def influencers_list(request):
    # You can add filtering logic here later
    influencers = Influencer.objects.all()
    return render(request, 'influencers.html', {'influencers': influencers})